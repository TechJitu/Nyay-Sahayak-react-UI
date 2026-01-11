import os
import base64
import io
import json
import shutil
import chromadb
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Request, Response
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from langchain_groq import ChatGroq
from groq import Groq
from fpdf import FPDF
from docx import Document
import PyPDF2
import requests
import time

# ==========================================
# API KEYS (Use Environment Variables)
# ==========================================
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")

if not GROQ_API_KEY:
    print("WARNING: GROQ_API_KEY not set! Please set it as an environment variable.")
    print("Run: set GROQ_API_KEY=your_api_key_here (Windows)")
    print("Or:  export GROQ_API_KEY=your_api_key_here (Linux/Mac)") 

app = FastAPI(title="Nyay Sahayak API", version="Final Hackathon Edition")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- DATABASE SETUP ---
try:
    chroma_client = chromadb.PersistentClient(path="./nyay_memory") 
    # Use ChromaDB's default embedding function (no heavy dependencies)
    vector_db = chroma_client.get_or_create_collection(name="legal_cases")
    print("Database Connected!")
except Exception as e:
    print(f"Database Error: {e}")
    vector_db = None

# --- AI MODEL SETUP (GROQ ONLY) ---
try:
    if GROQ_API_KEY:
        # 1. For Chat/Text Generation
        draft_llm = ChatGroq(groq_api_key=GROQ_API_KEY, model_name="llama-3.3-70b-versatile", temperature=0.3)
        vision_llm = ChatGroq(groq_api_key=GROQ_API_KEY, model_name="llama-3.2-11b-vision-preview", temperature=0)
        
        # 2. For Audio/Whisper (Direct Client)
        groq_client = Groq(api_key=GROQ_API_KEY)
        
        print("✅ Groq Models Ready!")
    else:
        draft_llm = None
        vision_llm = None
        groq_client = None
        print("Groq models not initialized - API key missing")
except Exception as e:
    print(f"Model Error: {e}")
    draft_llm = None
    vision_llm = None
    groq_client = None

# --- DATA MODELS ---
class UserQuery(BaseModel):
    question: str
    user_name: str = "User"
    role: str = "Citizen"
    language: str = "Hinglish"
    detail_level: str = "Detailed"
    state: str = "India (General)" 
    history: str = "" 

class ChatRequest(BaseModel):
    message: str
    history: str = ""

class RentAgreementQuery(BaseModel):
    landlord: str
    tenant: str
    rent: str
    address: str
    date: str

class NoticeRequest(BaseModel):
    voice_input: str

class ReportChatRequest(BaseModel):
    user_input: str 
    history: str = ""

# ==========================================
# LIVE STREAMING CHAT (GROQ LLAMA-3)
# ==========================================

async def generate_live_response(message, history):
    system_prompt = """You are Nyay Sahayak, an AI Legal Assistant for Indians. Provide accurate, helpful legal advice.

CRITICAL LANGUAGE RULES:
1. DETECT the language of the user's message.
2. Respond ENTIRELY in the SAME language - no mixing!
3. If English query → respond FULLY in English (use 1, 2, 3 for numbers)
4. If Hindi query → respond FULLY in Hindi (use १, २, ३ for numbers)
5. If Hinglish query → respond in Hinglish style

EXAMPLES:
- "How to file FIR?" → Full English response with 1, 2, 3
- "FIR कैसे file करें?" → Full Hindi response with १, २, ३
- "Mujhe help chahiye" → Hinglish response

RESPONSE FORMAT:
- Use **bold** for headings and important terms
- Use numbered lists for steps
- Structure with clear sections
- End with a helpful follow-up question

Keep responses professional, clear, and FULLY in the detected language."""
    full_prompt = f"{system_prompt}\n\nCONVERSATION HISTORY:\n{history}\n\nUSER: {message}\nAI:"
    
    try:
        if not draft_llm:
            yield "⚠️ Groq API Key missing or invalid."
            return

        async for chunk in draft_llm.astream(full_prompt):
            if chunk.content:
                yield chunk.content
    except Exception as e:
        yield f"Error: {str(e)}"

@app.post("/stream-chat")
async def stream_chat(request: ChatRequest):
    return StreamingResponse(
        generate_live_response(request.message, request.history), 
        media_type="text/plain"
    )

# ==========================================
# FILE REPORT INTERVIEW
# ==========================================
@app.post("/file-report-interview")
async def file_report_interview(data: ReportChatRequest):
    system_instruction = """
    ACT AS: An experienced, empathetic Police Officer (S.H.O) in India.
    GOAL: Gather details for an FIR (First Information Report).
    RULES: Ask ONLY ONE question at a time. Step-by-step.
    
    IMPORTANT LANGUAGE RULE: You MUST detect the language of the user's message and respond in THE SAME LANGUAGE.
    - If user speaks in English, respond ONLY in English.
    - If user speaks in Hindi, respond ONLY in Hindi.
    - If user speaks in Hinglish (mixed), respond in Hinglish.
    - Match the user's language exactly. Never switch languages unless the user does.
    
    Once all details are gathered, say "REPORT_COLLECTED".
    """
    full_prompt = f"{system_instruction}\nHISTORY:\n{data.history}\nUser: {data.user_input}\nAI:"
    
    try:
        if not draft_llm:
            return {"answer": "AI Offline"}
        res = draft_llm.invoke(full_prompt)
        return {"answer": res.content, "status": "active"}
    except Exception as e:
        return {"error": str(e)}

# ==========================================
# VOICE MESSAGE (GROQ WHISPER)
# ==========================================
@app.post("/voice-message")
async def voice_message(file: UploadFile = File(...), history: str = Form(default="")):
    print(f"Receiving Voice Note: {file.filename}")
    try:
        temp_filename = f"temp_{file.filename}"
        with open(temp_filename, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Transcribe using Groq Whisper
        if not groq_client:
             os.remove(temp_filename)
             return {"answer": "AI Config Error", "user_text": "Error"}

        with open(temp_filename, "rb") as audio_file:
            # ✅ UPDATED: Auto-Detect Language (Removed language='en')
            transcription = groq_client.audio.transcriptions.create(
                file=(temp_filename, audio_file.read()),
                model="whisper-large-v3", # Standard model
                response_format="json",
                # language="en",  <-- REMOVED THIS LINE TO ALLOW HINDI
                temperature=0.0
            )
        user_text = transcription.text
        
        # AI Response - Match user's language
        full_prompt = f"""ACT AS: Lawyer/Police. Keep it short and helpful.

IMPORTANT: Detect the language of the user's message and respond in THE SAME LANGUAGE.
- If user spoke in English, respond in English.
- If user spoke in Hindi, respond in Hindi.
- If user spoke in Hinglish, respond in Hinglish.
- Match their language exactly.

HISTORY:\n{history}\nUSER SAID: {user_text}"""
        res = draft_llm.invoke(full_prompt)
        ai_response = res.content

        os.remove(temp_filename)
        return {"user_text": user_text, "answer": ai_response}
    except Exception as e:
        print(f"Voice Error: {e}")
        return {"answer": "Error processing audio.", "user_text": "Error"}

# ==========================================
# DOC GENERATORS (PDF/DOCX)
# ==========================================
@app.post("/generate-legal-notice")
async def generate_legal_notice(data: NoticeRequest):
    extraction_prompt = f"""
    Extract details for Legal Notice. JSON ONLY.
    User Complaint: "{data.voice_input}"
    Keys: sender_name, receiver_name, amount, reason, act
    """
    try:
        res = draft_llm.invoke(extraction_prompt)
        content = res.content.replace("```json", "").replace("```", "").strip()
        details = json.loads(content)
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, txt="LEGAL NOTICE", ln=True, align='C')
        pdf.ln(10)
        pdf.set_font("Arial", size=12)
        notice_text = f"""To, {details.get('receiver_name', 'Receiver')}\n\nSubject: Legal Notice\n\nSir/Madam,\nOn behalf of {details.get('sender_name', 'Client')}, I state that you must resolve: {details.get('reason', 'Issue')}.\nOutstanding: {details.get('amount', 'N/A')}.\nFailure will lead to legal action under {details.get('act', 'Indian Laws')}."""
        pdf.multi_cell(0, 8, txt=notice_text)
        
        filename = f"Legal_Notice.pdf"
        pdf.output(filename)
        return FileResponse(path=filename, filename=filename, media_type='application/pdf')
    except Exception as e:
        return {"error": str(e)}

@app.post("/generate-rent-agreement")
async def generate_rent_agreement(data: RentAgreementQuery):
    try:
        doc = Document()
        doc.add_heading('RENT AGREEMENT', 0)
        doc.add_paragraph('Date: ' + data.date)
        doc.add_paragraph(f'LANDLORD: {data.landlord}', style='List Bullet')
        doc.add_paragraph(f'TENANT: {data.tenant}', style='List Bullet')
        filename = f"Rent_Agreement.docx"
        doc.save(filename)
        return FileResponse(path=filename, filename=filename, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return {"error": str(e)}

# Run the application (for Render deployment)
if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)

